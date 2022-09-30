from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

from .onto.leavedonto import OntoManager, LeavedOnto
from .utils import parse_vocab, parse_tagged_sentences

new = 'New Words'
legend = 'Legend'
shared = 'Shared Words'
cur = 'Current'
prev = 'Previous'
absent = 'Not in Current'
total_word_count = 'Total Word Count: {}'
vocab = 'Vocabulary'
FIELDS = ['word', 'origin']


def get_selected_fields(om, entry):
    filtered = []
    for f in FIELDS:
        value = om.onto1.get_field_value(entry, f)
        value = str(value)
        filtered.append(value)
    return filtered


def gen_vocab_report(onto_path, out_path, vocab_path, tagged_path):
    # total_data = gather_total_data(onto_path)
    # lessons_data = gather_lesson_data(onto_path)

    import pickle
    tmp = Path('tmp.pickle')
    if not tmp.is_file():
        words_data = gather_word_data(onto_path, vocab_path, tagged_path)
        pickle.dump(words_data, open(str(tmp), 'wb'))
    else:
        words_data = pickle.load(open(str(tmp), 'rb'))



    level = onto_path.stem

    # # format it in docx
    # total_file = out_path / f'{level} Vocab Report - Total.docx'
    # export_total_vocab_report(level, total_data, total_file)
    #
    # lessons_file = out_path / f'{level} Vocab Report - Lessons.docx'
    # export_lessons_vocab_report(level, lessons_data, lessons_file)

    words_file = out_path / f'{level} Words Report.docx'
    export_words_report(words_data, words_file)


def export_lessons_vocab_report(level, total_data, out_file):
    doc = Document()
    styles = doc.styles

    freq_style = styles.add_style('freq', WD_STYLE_TYPE.CHARACTER)
    freq_font = freq_style.font
    freq_font.name = 'Lato'
    freq_font.italic = True

    entry_style = styles.add_style('entry', WD_STYLE_TYPE.CHARACTER)
    entry_font = entry_style.font
    entry_font.size = Pt(7)
    entry_font.name = 'Lato Light'

    tree_style = styles.add_style('tree', WD_STYLE_TYPE.CHARACTER)
    tree_font = tree_style.font
    tree_font.name = 'Lato Light'
    tree_font.size = Pt(11)

    # LESSONS
    for section, data in total_data.items():
        doc.add_heading(section, 0)
        for subtitle, subdata in data.items():
            if subtitle == absent:
                doc.add_heading(subtitle.replace('Current', section), 1)
            else:
                doc.add_heading(subtitle, 1)

            if subtitle.startswith('Total'):
                par = doc.add_paragraph()
                for level in subdata:
                    line = '\t'.join(level)
                    run = par.add_run(line, style=tree_style)
                    if ':' in line:
                        run.bold = True
                    run.add_break()
            elif subtitle == new or subtitle == absent:
                for ssubtitle, ssubdata in subdata.items():
                    doc.add_heading(ssubtitle, 3)
                    par = doc.add_paragraph()
                    run = par.add_run(f'{len(ssubdata)} — ')
                    run.font.size = Pt(9)
                    for entry in ssubdata:
                        par.add_run(' '.join(entry), style=entry_style)
                        par.add_run(' ')
            elif subtitle == shared:
                for ssubtitle, ssubdata in subdata.items():
                    doc.add_heading(ssubtitle, 3)
                    par = doc.add_paragraph()
                    for sssubtitle, sssubdata in ssubdata.items():
                        run = par.add_run(sssubtitle + ': ', style=entry_style)
                        run.bold = True
                        par.add_run()
                        run = par.add_run(f'{len(sssubdata)} — ')
                        run.font.size = Pt(9)
                        for entry in sssubdata:
                            par.add_run(' '.join(entry), style=entry_style)
                            par.add_run(' ')
                        par.add_run().add_break()
                    par.runs[-1].text = ''
            else:
                raise ValueError('this is unexpected!')

    doc.save(out_file)


def export_total_vocab_report(level, total_data, out_file):
    doc = Document()
    styles = doc.styles

    freq_style = styles.add_style('freq', WD_STYLE_TYPE.CHARACTER)
    freq_font = freq_style.font
    freq_font.name = 'Lato'
    freq_font.italic = True

    entry_style = styles.add_style('entry', WD_STYLE_TYPE.CHARACTER)
    entry_font = entry_style.font
    entry_font.size = Pt(7)
    entry_font.name = 'Lato Light'

    tree_style = styles.add_style('tree', WD_STYLE_TYPE.CHARACTER)
    tree_font = tree_style.font
    tree_font.name = 'Lato Light'
    tree_font.size = Pt(11)

    # TITLE
    doc.add_heading(f'{level} Vocabulary Report - Total', 0)

    # TOTAL
    for section, data in total_data.items():
        doc.add_heading(section, 1)
        if section.startswith('Total'):
            par = doc.add_paragraph()
            for level in data:
                line = '\t'.join(level)
                run = par.add_run(line, style=tree_style)
                if ':' in line:
                    run.bold = True
                run.add_break()

        else:
            for subtitle, subdata in data.items():
                doc.add_heading(subtitle, 2)
                par = doc.add_paragraph()
                run = par.add_run(f'{len(subdata)} — ')
                run.font.size = Pt(9)
                for entry in subdata:
                    par.add_run(' '.join(entry), style=entry_style)
                    par.add_run(' ')

    doc.save(out_file)


def export_words_report(words_data, words_file):
    doc = Document()
    styles = doc.styles

    freq_style = styles.add_style('freq', WD_STYLE_TYPE.CHARACTER)
    freq_font = freq_style.font
    freq_font.name = 'Lato'
    freq_font.italic = True
    freq_font.size = Pt(7)

    entry_style = styles.add_style('entry', WD_STYLE_TYPE.CHARACTER)
    entry_font = entry_style.font
    entry_font.size = Pt(7)
    entry_font.name = 'Lato Light'

    tree_style = styles.add_style('tree', WD_STYLE_TYPE.CHARACTER)
    tree_font = tree_style.font
    tree_font.name = 'Lato Light'
    tree_font.size = Pt(5)

    for level, lessons in words_data.items():
        # TITLE
        doc.add_heading(f'Level {level}', 1)

        for lesson, words in lessons.items():
            doc.add_heading(f'Lesson {lesson}', 3)

            for word, occurences in words.items():
                word_heading = doc.add_heading(f'"{word}"', 5)
                sanity_freq = 0

                onto_par = doc.add_paragraph()
                onto_path = ''

                if 'sanity' in occurences:
                    sanity = occurences['sanity']
                    heading = doc.add_heading('Previous lessons', 7)
                    run = heading.runs[-1]
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    for f in sanity:
                        sanity_freq += f['freq']
                        par = doc.add_paragraph()
                        run = par.add_run(f'{f["origin"]}.docx, freq: {f["freq"]}, path: {"/".join(f["path"])}', style=freq_style)
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        if 'sentences' in f:
                            for sent in f['sentences']:
                                par.add_run('\n')
                                run = par.add_run(' '.join([s[0] for s in sent]), style=entry_style)
                                run.font.color.rgb = RGBColor(255, 0, 0)

                if not 'occurences' in occurences:
                    continue

                occurences = occurences['occurences']
                total_freq = occurences['freq'] + sanity_freq
                word_heading.add_run(f': {total_freq}')

                lessons_with_word = [o for o in occurences.keys() if o  != 'freq']
                for l_ww in lessons_with_word:
                    lesson_freq = occurences[l_ww]['freq']
                    doc.add_heading(f'{l_ww}: {lesson_freq}', 7)

                    par = doc.add_paragraph()

                    for f_ww in occurences[l_ww]['files']:

                        file_freq = f_ww['freq']
                        par.add_run(f'{file_freq} in ', style=freq_style)

                        f = f_ww['origin']
                        par.add_run(f'"{f}.docx"', style=freq_style)

                        path = '/'.join(f_ww['path'])
                        if not onto_path:
                            onto_path = path
                            onto_par.add_run(f'onto: {path}', style=freq_style)
                        elif path != onto_path:
                            par.add_run(f', onto: {path}', style=freq_style)

                        if 'sentences' in f_ww:
                            par.add_run('\n\t')
                            for sent in f_ww['sentences']:
                                par.add_run('{'+ ' '.join([s[0] for s in sent]) + '}, ', style=entry_style)
                            par.runs[-1].text = par.runs[-1].text.rstrip().rstrip(',')
                        par.add_run('\n')
                    par.runs[-1].text = par.runs[-1].text.rstrip()

    doc.save(words_file)


def gather_total_data(onto_path):
    om = OntoManager()
    for f in sorted(list(onto_path.glob('*.yaml'))):
        om.merge_to_onto(f)

    total_data = {}

    tree, total_words = om.onto1.export_tree_report()
    total_data[total_word_count.format(total_words)] = tree

    # gather data
    entries = om.onto1.ont.find_entries()

    # organize it for report
    total_data[vocab] = {}
    for p, e in entries:
        filtered = []
        for entry in e:
            filtered.append(get_selected_fields(om, entry))
        total_data[vocab]['/'.join(p)] = filtered

    return total_data


def gather_lesson_data(onto_path):
    # group ontos belonging to the same lesson
    levels = defaultdict(list)
    for f in sorted(list(onto_path.glob('*.yaml'))):
        lesson, _ = f.stem.split('-')
        levels[lesson].append(f)

    # generate one onto for each lesson
    all_ontos = []
    for l, ontos in levels.items():
        om = None
        for onto in ontos:
            ont = LeavedOnto(onto)
            if not om:
                om = OntoManager()
                om.onto1.ont_path = f.parent / (l + '_')  # required to merge these ontos

            om.merge_to_onto(ont, in_to_organize=False)
        all_ontos.append((l, om))

    previous_onto = OntoManager()
    report_data = {}
    for num, ao in enumerate(all_ontos):
        lesson, om = ao
        if lesson not in report_data:
            report_data[lesson] = {}

        # VOCAB TREE
        tree, total_words = om.onto1.export_tree_report()
        report_data[lesson][total_word_count.format(total_words)] = tree

        # WORD LISTS
        # everything is new vocab
        if num == 0:
            # gather data
            entries = om.onto1.ont.find_entries()

            # organize it for report
            report_data[lesson][new] = {}
            for p, e in entries:
                filtered_fields = []
                for entry in e:
                    filtered_fields.append(get_selected_fields(om, entry))
                report_data[lesson][new]['/'.join(p)] = filtered_fields

        # split words of lesson in: New, Shared, Unseen
        else:
            # gather data
            previous_onto.merge_to_onto(all_ontos[num-1][1].onto1, add_origin=False)
            current_only, common, previous_only = om.diff_ontos(previous_onto.onto1)

            # organize it for report
            report_data[lesson][new] = {}
            for p, e in current_only:
                title = '/'.join(p)
                if title not in report_data[lesson][new]:
                    report_data[lesson][new][title] = []

                report_data[lesson][new][title].append(get_selected_fields(om, e))

            report_data[lesson][shared] = {}
            for c, p in common:
                title = '/'.join(c[0])
                if title not in report_data[lesson][shared]:
                    report_data[lesson][shared][title] = {cur: [], prev: []}

                report_data[lesson][shared][title][cur].append(get_selected_fields(om, c[1]))
                report_data[lesson][shared][title][prev].append(get_selected_fields(om, p[1]))

            report_data[lesson][absent] = {}
            for p, e in previous_only:
                title = '/'.join(p)
                if title not in report_data[lesson][absent]:
                    report_data[lesson][absent][title] = []

                report_data[lesson][absent][title].append(get_selected_fields(om, e))

    return report_data


def gather_word_data(onto_path, vocab_path, tagged_path):
    word_data = process_n_filter_ontos(onto_path, vocab_path)
    retrieve_sentences(word_data, tagged_path)
    return word_data


def retrieve_sentences(word_data, tagged_path):
    sentences = parse_tagged_sentences(tagged_path)
    for level, lessons in word_data.items():
        for lesson, words in lessons.items():
            for word, occurences in words.items():
                if 'occurences' in occurences:
                    occurences = occurences['occurences']
                    for k, v in occurences.items():
                        if k != 'freq':
                            files = v['files']
                            for f in files:
                                origin = f['origin']
                                if 'sentences' in origin:
                                    pair = (word, f['POS'])
                                    for sent in sentences[origin]:
                                        if pair in sent:
                                            if 'sentences' not in f:
                                                f['sentences'] = []

                                            if sent not in f['sentences']:
                                                f['sentences'].append(sent)


def process_n_filter_ontos(onto_path, vocab_path):
    vocab = parse_vocab(vocab_path)

    om = OntoManager()
    for f in sorted(list(onto_path.glob('*.yaml'))):
        om.merge_to_onto(f)

    word_data = {}

    for level, lessons in vocab.items():
        if level not in word_data:
            word_data[level] = {}

        for lesson, a in lessons.items():
            if lesson not in word_data[level]:
                word_data[level][lesson] = {}

            field_type = a['legend'][2]
            words = a['words']
            for word, pos, field in words:
                if word not in word_data[level][lesson]:
                    word_data[level][lesson][word] = {}

                # 1. find word in onto
                results = om.onto1.find_word(word)

                # 2. filter results
                # 2.a. filter on pos
                if pos:
                    results = [(r, entry) for r, entry in results if r[0] == pos]
                # 2.b. filter on path-within-onto/categorisation
                if field_type == 'CAT':
                    if field:
                        results = [(path, entry) for path, entry in results if '/'.join(path) == field]
                # 2.c. filter on onto field
                elif field and field_type in om.onto1.ont.legend:
                    tmp = []
                    for path, entries in results:
                        tmp_res = []
                        for entry in entries:
                            if om.onto1.get_field_value(entry, field_type) == field:
                                tmp_res.append(entry)
                        if tmp_res:
                            tmp.append((path, tmp_res))
                    if tmp:
                        results = tmp
                else:
                    print(f"{field_type} is not a field in the ontology and can't be used to filter the results.")

                # 3. establish usage map of word
                # 3.1 sanity check: word has not been used in previous lessons
                for path, entries in results:
                    for entry in entries:
                        origins = om.onto1.get_field_value(entry, 'origin').split(' — ')
                        for o in origins:
                            filename, freq = o.split(':')
                            freq = int(freq)
                            o_lesson = filename.split('-')[0]
                            if o_lesson < lesson:
                                if 'sanity' not in word_data[level][lesson][word]:
                                    word_data[level][lesson][word]['sanity'] = []

                                word_data[level][lesson][word]['sanity'].append({'POS': pos, 'path': path, 'origin': filename, 'freq': freq})
                            else:
                                if 'occurences' not in word_data[level][lesson][word]:
                                    word_data[level][lesson][word]['occurences'] = {'freq': 0}

                                if o_lesson not in word_data[level][lesson][word]['occurences']:
                                    word_data[level][lesson][word]['occurences'][o_lesson] = {'freq': 0, 'files': []}

                                word_data[level][lesson][word]['occurences'][o_lesson]['files'].append({'POS': pos, 'path': path, 'origin': filename, 'freq': freq})
                                word_data[level][lesson][word]['occurences']['freq'] += freq
                                word_data[level][lesson][word]['occurences'][o_lesson]['freq'] += freq
    return word_data
